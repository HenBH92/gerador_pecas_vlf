import streamlit as st
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from manipulacao_arquivo import carregar_topicos_disponiveis, carregar_modelo, substituir_placeholders, adicionar_topicos, gerar_nome_arquivo, salvar_documento

CAMINHO_MODELO = r"C:\Users\hcost\Programas\PythonProject\Gerador de Peças VLF\base_defesa_inter.docx"
PASTA_TOPICOS = r"C:\Users\hcost\Programas\PythonProject\Gerador de Peças VLF\Topicos de Defesa"
CAMINHO_CONCLUSAO = r"C:\Users\hcost\Programas\PythonProject\Gerador de Peças VLF\base_conclusao_defesa_inter.docx"


def interface_streamlit():

    st.set_page_config(layout="wide")

    col_logo1, col_logo2, col_logo3 = st.columns([3, 4, 1])
    with col_logo2:
        st.image("logo_VLF.png", width=190)

    # Títulos
    st.markdown(
        '<h1 style="text-align: center; color: #201747">Gerador Automático de Peças VLF</h1>',
        unsafe_allow_html=True
    )
    st.markdown(
        '<h3 style="text-align: center; color: #201747">O programa está em desenvolvimento. A princípio estamos apenas com modelos de defesa do Inter em casos da Almaviva</h3>',
        unsafe_allow_html=True
    )

    # Lista Session State Tópics
    if "topicos_selecionados" not in st.session_state:
        st.session_state.topicos_selecionados = []

    # Campos de entrada para dados variáveis
    st.write("---")
    st.write("### Dados do Processo")
    vara_do_trabalho = st.text_input("Qual a Vara do Trabalho?").title()
    cidade = st.text_input("Qual a cidade?").title()
    estado = st.text_input("Qual o Estado (sigla)?", max_chars=2).upper()
    numero_processo = st.text_input("Qual o número do processo?")
    reclamante = st.text_input("Qual o nome do reclamante?").title()
    advogado = st.text_input("Quem é o advogado subscritor?").title()
    oab = st.text_input("Qual a OAB?")

    # Dicionário de dados inseridos
    dados_inseridos = {
        "[VARA DO TRABALHO]": vara_do_trabalho,
        "[CIDADE]": cidade,
        "[ESTADO]": estado,
        "[NUMERO DO PROCESSO]": numero_processo,
        "[RECLAMANTE]": reclamante,
        "[ADVOGADO]": advogado,
        "[Nº_OAB]": oab,
    }

    col1, col2 = st.columns([2, 1])

    # Seleção de tópicos existentes
    with col1:
        st.write("---")
        st.write("### Seleção de Tópicos")
        topicos_disponiveis = carregar_topicos_disponiveis(PASTA_TOPICOS)
        for topico in topicos_disponiveis:
            if st.checkbox(topico):
                if topico not in st.session_state.topicos_selecionados:
                    st.session_state.topicos_selecionados.append(topico)
            else:
                if topico in st.session_state.topicos_selecionados:
                    st.session_state.topicos_selecionados.remove(topico)

    # Quadro exibindo os tópicos:
    with col2:
        st.write("---")
        st.write("### Ordem dos Tópicos na Defesa")
        if st.session_state.topicos_selecionados == []:
            st.write("Selecione um tópico")
        else:
            for index, item in enumerate(st.session_state.topicos_selecionados):
                st.write(f"{index + 1} - {item}")


    # Novos tópicos com session_state
    st.write("---")
    st.write("### Novos Tópicos")
    if "novos_topicos" not in st.session_state:
        st.session_state.novos_topicos = []

    escolha = st.radio("Você deseja incluir algum tópico que não está na lista?", ["Sim", "Não"]).capitalize()
    if escolha == "Sim":
        nome_novo_topico = st.text_input("Qual o título do tópico?")
        if st.button("Adicionar novo tópico"):
            if nome_novo_topico and nome_novo_topico not in st.session_state.novos_topicos:
                st.session_state.novos_topicos.append(nome_novo_topico)
                st.success(f"Tópico {nome_novo_topico} adicionado!")
            elif nome_novo_topico in st.session_state.novos_topicos:
                st.warning("Este tópico já foi adicionado anteriormente!")
            else:
                st.error("Por favor, insira um título válido para o tópico.")

    # Exibir tópicos adicionados (opcional, para feedback)
    if st.session_state.novos_topicos:
        st.write(f"Tópicos adicionados: {st.session_state.novos_topicos}")


    # Seleção do cliente
    st.write("---")
    st.write("### Cliente")
    clientes = {"Banco Inter": "INTER"}
    cliente = st.selectbox("Selecione o cliente", list(clientes.keys()))
    codigo_cliente = clientes[cliente]



    # Botão para gerar o documento
    st.write("---")
    if st.button("Gerar Documento"):
        doc = carregar_modelo(CAMINHO_MODELO)
        adicionar_topicos(doc, st.session_state.topicos_selecionados, PASTA_TOPICOS, st.session_state.novos_topicos)
        conclusao_doc = carregar_modelo(CAMINHO_CONCLUSAO)
        for paragrafo in conclusao_doc.paragraphs:
            doc.add_paragraph(paragrafo.text, style=paragrafo.style)
        if conclusao_doc.tables:
            tabela = conclusao_doc.tables[0]
            linhas = len(tabela.rows)
            colunas = len(tabela.columns)
            nova_tabela = doc.add_table(linhas, colunas)
            for linha in range(linhas):
                for coluna in range (colunas):
                    celula_original = tabela.cell(linha, coluna)
                    celula_nova = nova_tabela.cell(linha, coluna)
                    for paragrafo in celula_original.paragraphs:
                        novo_paragrafo = celula_nova.add_paragraph()
                        novo_paragrafo.alignment = paragrafo.alignment
                        for run in paragrafo.runs:
                            novo_run = novo_paragrafo.add_run(run.text)
                            novo_run.bold = run.bold
                            novo_run.italic = run.italic
                            novo_run.font.name = run.font.name
                            novo_run.font.size = run.font.size
        substituir_placeholders(doc, dados_inseridos)
        nome_arquivo = gerar_nome_arquivo(reclamante, codigo_cliente)
        arquivo_bytes = salvar_documento(doc)
        st.download_button(
            label="Baixar Documento",
            data=arquivo_bytes,
            file_name=f"{nome_arquivo}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.success("Documento gerado! Clique para baixar.")

interface_streamlit()
