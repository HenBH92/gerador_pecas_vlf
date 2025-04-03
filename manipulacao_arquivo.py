import os
import glob
import datetime
import os.path
import docx
from io import BytesIO


def carregar_topicos_disponiveis(pasta_topicos):
  topicos_disponiveis = []
  caminho_busca = os.path.join(pasta_topicos, "*.docx")
  topicos = glob.glob(caminho_busca)
  for topico in topicos:
    nome_arquivo = os.path.basename(topico)
    nome_sem_extensao = (os.path.splitext(nome_arquivo)[0]).title()
    topicos_disponiveis.append(nome_sem_extensao)
  return topicos_disponiveis


def carregar_modelo(caminho_modelo):
  """Carrega o modelo Word base para a contestação.
Args:
    caminho_modelo (str): Caminho completo do arquivo .docx do modelo.
Returns:
    Document: Objeto do documento Word carregado.
"""
  modelo_base = docx.Document(caminho_modelo)
  return modelo_base


def substituir_placeholders(doc, dados):
  for paragrafo in doc.paragraphs:
    for placeholder, valor in dados.items():
      if placeholder in paragrafo.text:
        paragrafo.text = paragrafo.text.replace(placeholder, valor)
  if doc.tables:
    for tabela in doc.tables:
      for linha in range(len(tabela.rows)):
        for coluna in range(len(tabela.columns)):
          celula = tabela.cell(linha, coluna)
          for paragrafo in celula.paragraphs:
            for run in paragrafo.runs:
              for placeholder, valor in dados.items():
                if placeholder in run.text:
                  run.text = run.text.replace(placeholder, valor)





def adicionar_topicos(doc, topicos_selecionados, pasta_topicos, novos_topicos=None):
  for topico in topicos_selecionados:
    caminho = os.path.join(pasta_topicos, f'{topico}.docx')
    topico_doc = docx.Document(caminho)
    for paragrafo in topico_doc.paragraphs:
      doc.add_paragraph(paragrafo.text, style = paragrafo.style)
  if novos_topicos:
    print("Estilos disponíveis no doc:", [style.name for style in doc.styles])
    for nome in novos_topicos:
      doc.add_heading(nome, level=2)
      doc.add_paragraph("Insira suas razões", style="AParágrafo")


def gerar_nome_arquivo(nome_reclamante, codigo_cliente):
  """Gera o nome do arquivo no formato 'ANO-MÊS-DIA - CÓDIGO DO CLIENTE - Contestação (NOME DO RECLAMANTE)'.
Args:
    nome_reclamante (str): Nome do reclamante.
    codigo_cliente (str): Código do cliente.
Returns:
    str: Nome formatado do arquivo.
"""
  data = datetime.datetime.now()
  data_formatada = data.strftime("%Y-%m-%d")
  nome_arquivo = f"{data_formatada} - {codigo_cliente} - Contestação ({nome_reclamante})"
  return nome_arquivo


def salvar_documento(doc):
  """Retorna o documento como bytes para download.
  Args:
      doc (Document): Objeto do documento Word.
  Returns:
      bytes: Conteúdo do arquivo em formato binário.
  """
  buffer = BytesIO()
  doc.save(buffer)
  buffer.seek(0)
  return buffer.getvalue()